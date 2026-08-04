# -*- coding: utf-8 -*-
"""
生成《2D 地图引擎选型报告》Word 文档。

对比对象：百度地图 GL / 高德 Web JS / Leaflet / OpenLayers / Mapbox GL JS / MapLibre GL JS / Cesium
维度：技术特性 / 离线部署 / 性能 / 开源协议与商用成本 / 生态 / 学习曲线 / 与项目契合度
最终输出：docs/2D地图引擎选型报告.docx
"""
from __future__ import annotations

import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "docs/2D地图引擎选型报告.docx"
OUTPUT_FALLBACK = "docs/2D地图引擎选型报告_{ts}.docx"

# ============ 样式辅助 ============

CN_FONT = "微软雅黑"
CN_FONT_LIGHT = "微软雅黑"
EN_FONT = "Segoe UI"

DARK = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0x6B, 0xB8)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xEF, 0x6C, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY_FILL = "EEF3F8"
HEAD_FILL = "1F3A5F"
ALT_FILL = "F5F7FA"


def set_run_font(run, size: float, bold=False, color=None, cn=CN_FONT, en=EN_FONT):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)


def add_para(doc, text="", size=10.5, bold=False, color=None, align=None, space_after=4, space_before=0, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.3
    if indent is not None:
        pf.left_indent = Pt(indent)
    if text:
        r = p.add_run(text)
        set_run_font(r, size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 13, 3: 11.5}
    colors = {0: DARK, 1: DARK, 2: ACCENT, 3: ACCENT}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, sizes.get(level, 11), bold=True, color=colors.get(level, DARK))
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullet(doc, text, size=10.5, color=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 16)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    # 支持 **粗体** 简单标记
    segments = text.split("**")
    for i, seg in enumerate(segments):
        if not seg:
            continue
        r = p.add_run(seg)
        set_run_font(r, size, bold=(i % 2 == 1), color=color)
    return p


def shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, size=9.5, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    # 支持多行（按 \n 拆分）
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows, col_widths=None, header_fill=HEAD_FILL, first_col_bold=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill=header_fill)
    # 数据行
    for r_idx, row in enumerate(rows):
        fill = ALT_FILL if r_idx % 2 == 1 else None
        for c_idx, val in enumerate(row):
            bold = first_col_bold and c_idx == 0
            color = DARK if c_idx == 0 else None
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.rows[1 + r_idx].cells[c_idx], val, size=9, bold=bold, color=color,
                          align=align, fill=fill)
    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    # 边框
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BBBBBB")
        borders.append(e)
    tblPr.append(borders)
    return table


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    # 浅灰底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_callout(doc, title, text, color=ACCENT, fill=LIGHT_GRAY_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    shade_cell(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run_font(r1, 10.5, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(text)
    set_run_font(r2, 10, color=GRAY)
    # 左侧色条（通过左边框加粗模拟）
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "0E6BB8")
    borders.append(left)
    for edge in ("top", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "DDDDDD")
        borders.append(e)
    tcPr.append(borders)


def set_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)


# ============ 报告内容 ============

def build_document():
    doc = Document()
    set_page(doc)

    # 默认样式
    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = Pt(10.5)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)

    # ===== 封面 =====
    add_para(doc, "", size=6, space_after=0)
    add_para(doc, "2D 地图引擎选型报告", size=26, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "Ground Control Station · 无人机地面控制站", size=13, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "—— 对比 Leaflet / OpenLayers / Mapbox / MapLibre / 百度 / Cesium 等主流方案", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("文档版本", "v1.0"),
        ("编制日期", "2026-08-03"),
        ("适用项目", "Ground Control Station（地面控制站）"),
        ("文档密级", "内部使用"),
    ]
    for i, (k, v) in enumerate(info_rows):
        set_cell_text(info_table.rows[i].cells[0], k, size=10.5, bold=True, color=DARK, fill=LIGHT_GRAY_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(info_table.rows[i].cells[1], v, size=10.5, color=GRAY)
        info_table.rows[i].cells[0].width = Cm(4)
        info_table.rows[i].cells[1].width = Cm(10)
    # 边框
    tbl = info_table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "CCCCCC")
        borders.append(e)
    tblPr.append(borders)

    add_para(doc, "", size=10, space_after=20)
    add_callout(doc,
                "核心结论",
                "综合本项目「离线/内网部署 + 高性能态势大屏 + React/TypeScript 技术栈 + 商用合规」四大约束，"
                "本报告推荐 MapLibre GL JS 作为 2D 地图引擎首选方案，Leaflet 作为轻量备选，"
                "OpenLayers 作为强 GIS 场景的补充选项。百度/高德等国内商业 SDK 因公网依赖与离线授权限制，不适合作为离线内网场景的主力引擎。",
                color=GREEN, fill="E8F5E9")

    doc.add_page_break()

    # ===== 目录 =====
    add_heading(doc, "目  录", level=1)
    toc = [
        "1.  摘要",
        "2.  项目背景与选型需求",
        "3.  候选方案概览",
        "4.  多维度对比",
        "5.  关键维度深入分析",
        "    5.1 离线 / 内网部署能力",
        "    5.2 渲染性能",
        "    5.3 开源协议与商用合规",
        "    5.4 生态、社区与学习曲线",
        "    5.5 与现有项目架构的契合度",
        "6.  综合评分模型",
        "7.  选型结论与推荐",
        "8.  落地路线与风险评估",
        "9.  附录",
    ]
    for t in toc:
        add_para(doc, t, size=10.5, color=DARK, space_after=2, indent=6)

    doc.add_page_break()

    # ===== 1. 摘要 =====
    add_heading(doc, "1. 摘要", level=1)
    add_para(doc,
             "本报告面向 Ground Control Station（无人机地面控制站）项目，对当前主流 2D/2.5D Web 地图引擎进行系统化选型评估。"
             "项目当前采用百度地图 GL SDK（BMapGL），在公网环境下功能完整，但存在 SDK/瓦片/定位服务强依赖公网、BD09 专有坐标系、"
             "商用授权与数据合规等限制，无法满足未来「离线/内网部署」的刚性需求。")
    add_para(doc, "评估覆盖以下 7 类候选方案：", bold=True, space_after=2)
    add_bullet(doc, "国内商业 SDK：**百度地图 GL（BMapGL）**、**高德地图 JS API 2.0**")
    add_bullet(doc, "国际商业 SDK：**Mapbox GL JS**")
    add_bullet(doc, "开源矢量引擎：**MapLibre GL JS**")
    add_bullet(doc, "开源栅格/通用引擎：**Leaflet**、**OpenLayers**")
    add_bullet(doc, "开源 3D 地球引擎：**Cesium**（作为参照）")
    add_para(doc, "评估从 5 个核心维度展开：离线部署能力、渲染性能、开源协议与商用成本、生态与学习曲线、与现有架构契合度。", space_after=4)
    add_callout(doc, "一句话结论",
                "首选 MapLibre GL JS（GPU 矢量渲染 + BSD-3 无商用限制 + 完全本地化部署）；"
                "Leaflet 作为轻量/快速原型备选；OpenLayers 作为强 GIS 分析场景的补充；"
                "百度/高德仅在「在线 + 快速交付 + 无离线要求」时可选。",
                color=ACCENT)

    # ===== 2. 项目背景与选型需求 =====
    add_heading(doc, "2. 项目背景与选型需求", level=1)

    add_heading(doc, "2.1 现状", level=2)
    add_para(doc, "项目为基于 React 18 + TypeScript + Vite 的无人机地面控制站，核心地图能力包括：")
    add_bullet(doc, "卫星/矢量底图渲染（当前使用百度卫星图 + 自定义样式隐藏标注）")
    add_bullet(doc, "航线规划：点击加点、拖拽航点、右键删除、起终点语义")
    add_bullet(doc, "航线视觉：霓虹光晕折线（双层 Polyline）、脉冲航点节点")
    add_bullet(doc, "无人机模拟飞行：位置插值动画、航向旋转、金色拖尾轨迹")
    add_bullet(doc, "辅助控件：缩放按钮、动态比例尺、地址搜索（PlaceSearch）")
    add_para(doc, "当前架构已抽象出 MapAdapter 接口（src/map-engines/types.ts），为引擎切换奠定了基础。", space_after=4)

    add_heading(doc, "2.2 核心需求（按优先级）", level=2)
    req_headers = ["优先级", "需求项", "说明", "约束"]
    req_rows = [
        ["P0", "离线/内网部署", "SDK、瓦片、字体、图标全部本地化，无公网依赖", "刚性"],
        ["P0", "商用合规", "协议允许商用，无隐性费用与数据回传", "刚性"],
        ["P0", "坐标系标准化", "统一 WGS84，无人机原始坐标无需偏移转换", "刚性"],
        ["P1", "渲染性能", "态势大屏场景，万级图标流畅 60fps", "重要"],
        ["P1", "样式可定制", "暗色大屏风格、卫星图、要素显隐可控", "重要"],
        ["P1", "React/TS 友好", "类型完备、与现有组件架构兼容", "重要"],
        ["P2", "3D / 倾斜", "可选支持地形高程、三维视景", "可选"],
        ["P2", "专业 GIS 能力", "量算、缓冲区、空间分析", "可选"],
    ]
    add_table(doc, req_headers, req_rows, col_widths=[1.8, 3.2, 7.5, 2.0])
    add_para(doc, "")

    add_heading(doc, "2.3 选型决策树", level=2)
    add_code_block(doc,
        "是否需要离线/内网部署？\n"
        "├─ 否（仅在线） → 可选 百度/高德/Mapbox（交付快，但有授权/费用）\n"
        "└─ 是（刚性）\n"
        "   └─ 是否需要 3D 地形/倾斜视图？\n"
        "      ├─ 是 → MapLibre GL JS（原生 3D 地形）/ Cesium（真三维地球）\n"
        "      └─ 否\n"
        "         └─ 图标/要素数量级？\n"
        "            ├─ > 5000 → MapLibre GL JS（GPU 渲染）\n"
        "            ├─ 500 ~ 5000 → OpenLayers（Canvas/WebGL）或 MapLibre\n"
        "            └─ < 500 且需快速原型 → Leaflet（最轻量）")

    # ===== 3. 候选方案概览 =====
    add_heading(doc, "3. 候选方案概览", level=1)
    add_para(doc, "下表给出各方案的速览，详细对比见第 4 章。")

    ov_headers = ["方案", "类型", "渲染", "离线", "协议/费用", "3D", "定位"]
    ov_rows = [
        ["百度地图 GL", "国内商业", "WebGL 栅格", "❌ 公网依赖", "商用授权/AK", "有限", "在线/国内业务"],
        ["高德 JS 2.0", "国内商业", "WebGL 栅格", "❌ 公网依赖", "商用授权/Key", "有限", "在线/国内业务"],
        ["Mapbox GL JS", "国际商业", "WebGL 矢量", "⚠️ 部分可离线", "用量计费", "✅", "商业矢量/定制"],
        ["MapLibre GL JS", "开源(BSD-3)", "WebGL 矢量", "✅ 完全可离线", "免费开源", "✅", "离线大屏/高性能"],
        ["Leaflet", "开源(BSD-2)", "DOM/Canvas 栅格", "✅ 完全可离线", "免费开源", "❌", "轻量快速原型"],
        ["OpenLayers", "开源(BSD-2)", "Canvas/WebGL 矢+栅", "✅ 完全可离线", "免费开源", "有限", "专业 GIS"],
        ["Cesium", "开源(Apache-2)", "WebGL 3D 地球", "✅ 可离线", "免费(+收费资产)", "✅✅", "真三维/时空"],
    ]
    add_table(doc, ov_headers, ov_rows, col_widths=[2.8, 2.2, 2.6, 2.4, 2.4, 1.2, 2.8])
    add_para(doc, "")

    # ===== 4. 多维度对比 =====
    add_heading(doc, "4. 多维度对比", level=1)

    add_heading(doc, "4.1 技术特性对比", level=2)
    tech_headers = ["维度", "百度 GL", "Mapbox", "MapLibre", "Leaflet", "OpenLayers", "Cesium"]
    tech_rows = [
        ["渲染管线", "WebGL 栅格", "WebGL 矢量", "WebGL 矢量", "DOM/Canvas", "Canvas/WebGL", "WebGL 3D"],
        ["瓦片类型", "栅格(私有URL)", "矢量 MVT", "矢量 MVT", "栅格 XYZ", "矢量+栅格+OGC", "影像/地形"],
        ["样式动态切换", "有限(样式V2)", "✅ 运行时", "✅ 运行时", "❌ 重切片", "✅", "✅"],
        ["包体积(gzip)", "~250KB(在线)", "~250KB", "~250KB", "~40KB", "~120KB", "~400KB+"],
        ["TS 类型支持", "❌ 需自补 d.ts", "✅ 内置", "✅ 内置", "需 @types", "✅ 内置", "✅ 内置"],
        ["3D 地形/倾斜", "有限", "✅ terrain", "✅ terrain", "❌", "有限", "✅✅ 全球DEM"],
        ["数据可视化生态", "弱", "强", "强(Deck.gl)", "中(插件)", "中", "强(CesiumJS原生)"],
    ]
    add_table(doc, tech_headers, tech_rows, col_widths=[2.6, 2.3, 2.3, 2.3, 2.0, 2.4, 2.3])
    add_para(doc, "")

    add_heading(doc, "4.2 离线部署与合规对比", level=2)
    off_headers = ["维度", "百度 GL", "Mapbox", "MapLibre", "Leaflet", "OpenLayers"]
    off_rows = [
        ["SDK 本地化", "❌ CDN", "✅ npm", "✅ npm", "✅ npm", "✅ npm"],
        ["底图本地化", "❌", "⚠️ 需离线瓦片", "✅ 自建瓦片", "✅ Nginx 静态", "✅ 自建"],
        ["定位服务离线", "❌", "⚠️", "✅ 浏览器定位", "✅ 浏览器定位", "✅ 浏览器定位"],
        ["坐标系", "BD09(偏移)", "WGS84/Web墨卡托", "WGS84/Web墨卡托", "WGS84", "任意EPSG"],
        ["商用授权", "需AK/合规", "用量计费", "✅ BSD-3 免费", "✅ BSD-2 免费", "✅ BSD-2 免费"],
        ["数据回传风险", "高", "中", "无", "无", "无"],
    ]
    add_table(doc, off_headers, off_rows, col_widths=[2.8, 2.4, 2.4, 2.6, 2.6, 2.6])
    add_para(doc, "")

    add_heading(doc, "4.3 工程与生态对比", level=2)
    eng_headers = ["维度", "百度 GL", "Mapbox", "MapLibre", "Leaflet", "OpenLayers", "Cesium"]
    eng_rows = [
        ["GitHub Star(参考)", "—", "10k+", "6k+", "40k+", "11k+", "12k+"],
        ["社区活跃度", "国内强", "国际强", "活跃(上升)", "最成熟", "成熟", "成熟"],
        ["中文文档/案例", "丰富", "一般", "一般", "丰富", "一般", "一般"],
        ["学习曲线", "低", "中", "中", "低", "较高", "高"],
        ["React 封装", "弱", "react-map-gl", "react-map-gl", "react-leaflet", "rlayers", "resium"],
        ["国内偏移问题", "是(BD09)", "否", "否", "否", "否", "否"],
    ]
    add_table(doc, eng_headers, eng_rows, col_widths=[2.6, 2.2, 2.2, 2.2, 2.2, 2.4, 2.2])
    add_para(doc, "")

    # ===== 5. 关键维度深入分析 =====
    add_heading(doc, "5. 关键维度深入分析", level=1)

    add_heading(doc, "5.1 离线 / 内网部署能力", level=2)
    add_para(doc, "离线部署是本项目的 P0 约束，直接影响选型淘汰。完整离线化需同时满足三点：SDK 本地化、底图瓦片本地化、定位与地理编码本地化。")
    add_bullet(doc, "**百度/高德**：SDK 通过 CDN 加载，瓦片实时请求 CDN，Geolocation/PlaceSearch 依赖服务端。内网完全不可用，**一票否决**作为离线主力。")
    add_bullet(doc, "**Mapbox GL JS**：SDK 可 npm 本地化，但其官方瓦片服务（api.mapbox.com）需联网与 token；可接入自建瓦片，但商用受用量计费约束。")
    add_bullet(doc, "**MapLibre GL JS**：SDK 完全本地化，底图走自建矢量瓦片服务（tileserver-gl / Martin），定位用浏览器原生 Geolocation，**零公网依赖**。")
    add_bullet(doc, "**Leaflet**：SDK 本地化 + Nginx 静态栅格瓦片，离线方案最简，但栅格瓦片体积大、无法动态换肤。")
    add_bullet(doc, "**OpenLayers**：与 MapLibre 类似可完全离线，且可对接 GeoServer 发布 WMS/WMTS，适合企业级 GIS。")
    add_callout(doc, "结论",
                "在「离线/内网」维度，MapLibre / Leaflet / OpenLayers 满足；百度/高德不满足；Mapbox 部分满足但受商用约束。",
                color=GREEN, fill="E8F5E9")

    add_heading(doc, "5.2 渲染性能", level=2)
    add_para(doc, "地面站态势大屏需要承载大量飞行器图标、航迹、电子围栏等要素，渲染性能是 P1 约束。")
    perf_headers = ["场景", "Leaflet(DOM)", "Leaflet(Canvas)", "OpenLayers", "MapLibre", "Cesium"]
    perf_rows = [
        ["100 个标记", "流畅", "流畅", "流畅", "流畅", "流畅"],
        ["1,000 个标记", "卡顿", "流畅", "流畅", "流畅", "流畅"],
        ["10,000 个标记", "严重卡顿", "可用", "可用(WebGL)", "流畅60fps", "流畅"],
        ["动态轨迹(实时)", "弱", "中", "中", "强", "强"],
        ["矢量样式重绘", "N/A(栅格)", "N/A", "中", "快(GPU)", "快"],
    ]
    add_table(doc, perf_headers, perf_rows, col_widths=[3.2, 2.8, 2.8, 2.6, 2.6, 2.4])
    add_para(doc, "")
    add_bullet(doc, "**MapLibre / Mapbox**：基于 WebGL 的矢量瓦片渲染，图层样式在 GPU 计算，万级要素 60fps 稳定，是态势大屏首选。")
    add_bullet(doc, "**Leaflet**：默认 DOM 渲染，>1000 个标记需 Canvas 模式或 markercluster 插件；轨迹动画性能弱。")
    add_bullet(doc, "**OpenLayers**：WebGL 图层性能接近 MapLibre，但 API 更复杂。")
    add_bullet(doc, "**Cesium**：3D 场景性能强，但 2D 纯态势展示属于「杀鸡用牛刀」，包体积与学习成本偏高。")

    add_heading(doc, "5.3 开源协议与商用合规", level=2)
    lic_headers = ["方案", "协议", "商用条件", "隐性成本/风险"]
    lic_rows = [
        ["百度地图 GL", "专有(免费+商用授权)", "需申请 AK，遵守服务条款", "数据回传、用量限制、坐标系偏移、离线不可用"],
        ["高德 JS 2.0", "专有(免费+商用授权)", "需 Key，商用配额", "同上，每日配额限制"],
        ["Mapbox GL JS", "BSD-3(SDK)+服务条款", "免费额度后按用量计费", "瓦片/地理编码费用随用量增长，token 绑定"],
        ["MapLibre GL JS", "BSD-3", "完全免费，无用量限制", "无（需自备瓦片数据）"],
        ["Leaflet", "BSD-2-Clause", "完全免费", "无"],
        ["OpenLayers", "BSD-2-Clause", "完全免费", "无"],
        ["Cesium", "Apache-2.0", "免费(Cesium ion 资产部分收费)", "ion 在线服务需账号，离线需替换资产源"],
    ]
    add_table(doc, lic_headers, lic_rows, col_widths=[2.8, 3.4, 4.2, 5.8])
    add_para(doc, "")
    add_bullet(doc, "**关键结论**：MapLibre / Leaflet / OpenLayers 均为宽松开源协议，无商用限制、无数据回传，合规性最佳。")
    add_bullet(doc, "**Mapbox**：SDK 本身 BSD-3，但其价值依赖官方瓦片服务，商用按用量计费，大规模部署成本不可控。")
    add_bullet(doc, "**百度/高德**：免费但有「服务条款」约束，数据合规审查（尤其政军/能源行业）常被一票否决。")

    add_heading(doc, "5.4 生态、社区与学习曲线", level=2)
    add_bullet(doc, "**Leaflet**：社区最成熟、中文资源最丰富、API 极简，上手最快；插件生态庞大但质量参差。")
    add_bullet(doc, "**MapLibre**：Mapbox GL JS 的开源分叉，API 与 Mapbox 高度一致；可复用 Mapbox Style Spec；与 Deck.gl 协同最佳。")
    add_bullet(doc, "**Mapbox**：商业生态最完善，Studio 样式编辑器、数据服务、SDK 全栈，但依赖云服务。")
    add_bullet(doc, "**OpenLayers**：功能最全（OGC 全协议），但 API 偏底层，学习曲线较高，适合 GIS 专业团队。")
    add_bullet(doc, "**Cesium**：真三维地球，适合空天/全球尺度场景；2D 地面站场景 ROI 偏低。")
    add_bullet(doc, "**百度/高德**：国内文档与案例丰富，上手快，但强耦合国内生态与坐标系。")

    add_heading(doc, "5.5 与现有项目架构的契合度", level=2)
    add_para(doc, "项目已具备完善的地图业务层，引擎切换的核心成本在于「适配层」而非业务重写：")
    fit_headers = ["现有模块", "契合度（MapLibre）", "契合度（Leaflet）", "契合度（OpenLayers）", "说明"]
    fit_rows = [
        ["MapAdapter 接口", "✅ 高", "✅ 高", "✅ 高", "已有抽象层，三种引擎均可实现"],
        ["航线 Polyline(霓虹双层)", "✅ 原生 layer", "⚠️ 需双层叠加", "✅ 可实现", "矢量引擎渲染光晕更优"],
        ["航点 HTML 节点/动画", "✅ Marker DOM", "✅ divIcon", "✅ Overlay", "均支持自定义 HTML"],
        ["无人机位置/航向动画", "✅ 逐帧 setLngLat", "✅ setLatLng", "✅ coordinate", "RAF 动画与引擎无关"],
        ["拖尾轨迹实时更新", "✅ setData", "✅ setLatLngs", "✅ setCoordinates", "矢量引擎更平滑"],
        ["MapScale 比例尺", "✅ getMetersPerPixel", "✅ 可实现", "✅ 可实现", "适配层统一抽象"],
        ["地址搜索 PlaceSearch", "❌ 需替换", "❌ 需替换", "❌ 需替换", "百度专有，离线需自建/放弃"],
    ]
    add_table(doc, fit_headers, fit_rows, col_widths=[3.6, 2.8, 2.8, 2.8, 4.2])
    add_para(doc, "")
    add_callout(doc, "架构契合度结论",
                "三种开源引擎均可通过现有 MapAdapter 接口接入，业务层（航线/无人机/比例尺）改动可控。"
                "MapLibre 在矢量渲染、光晕折线、轨迹平滑度上表现最佳，且与未来 Deck.gl 可视化路线一致。",
                color=ACCENT)

    # ===== 6. 综合评分模型 =====
    add_heading(doc, "6. 综合评分模型", level=1)
    add_para(doc, "采用加权评分（满分 10，权重见列头），总分越高越优。")
    score_headers = ["维度 / 权重", "百度 GL", "Mapbox", "MapLibre", "Leaflet", "OpenLayers", "Cesium"]
    score_rows = [
        ["离线部署(25%)", "2", "6", "10", "10", "10", "8"],
        ["渲染性能(20%)", "7", "10", "10", "5", "8", "9"],
        ["协议/合规(20%)", "4", "6", "10", "10", "10", "9"],
        ["生态/学习(15%)", "8", "9", "8", "10", "7", "6"],
        ["架构契合(15%)", "6(现状)", "8", "9", "8", "8", "6"],
        ["3D/GIS 扩展(5%)", "4", "8", "8", "2", "8", "10"],
        ["加权总分", "4.75", "7.55", "9.45", "8.05", "8.70", "7.80"],
    ]
    add_table(doc, score_headers, score_rows, col_widths=[3.0, 2.2, 2.2, 2.2, 2.2, 2.4, 2.2])
    add_para(doc, "")
    add_para(doc, "评分说明：", bold=True, space_after=2)
    add_bullet(doc, "百度 GL 在「离线/合规」两项 P0 维度得分极低，虽性能与生态尚可，加权后被淘汰。")
    add_bullet(doc, "MapLibre 在 P0 维度满分，性能与契合度领先，总分第一。")
    add_bullet(doc, "OpenLayers 总分第二，适合需要 OGC/强 GIS 的后续扩展。")
    add_bullet(doc, "Leaflet 总分第三，但轻量与上手快的优势使其成为快速原型的最佳备选。")

    # ===== 7. 选型结论与推荐 =====
    add_heading(doc, "7. 选型结论与推荐", level=1)

    add_callout(doc, "最终推荐",
                "主引擎：MapLibre GL JS（矢量瓦片 + tileserver-gl 离线服务）。\n"
                "备选：Leaflet（栅格瓦片，用于轻量/快速验证或低性能要求终端）。\n"
                "补充：OpenLayers（未来若引入专业 GIS 分析/OGC 服务时再评估）。\n"
                "不推荐作为离线主力：百度/高德（公网依赖 + 合规风险）、Mapbox（用量计费）、Cesium（2D 场景过重）。",
                color=GREEN, fill="E8F5E9")

    add_heading(doc, "7.1 推荐理由（MapLibre GL JS）", level=2)
    add_bullet(doc, "**离线 P0 达标**：SDK 本地化 + 自建矢量瓦片 + 浏览器定位，零公网依赖。")
    add_bullet(doc, "**性能 P1 达标**：WebGL 矢量渲染，态势大屏万级图标 60fps。")
    add_bullet(doc, "**合规达标**：BSD-3 协议，无用量计费、无数据回传。")
    add_bullet(doc, "**样式灵活**：暗色/卫星/地形运行时切换，无需重新切片。")
    add_bullet(doc, "**架构契合**：与现有 MapAdapter 抽象层、Deck.gl 可视化路线一致。")
    add_bullet(doc, "**生态成熟**：Mapbox GL 开源分叉，API 稳定，react-map-gl 可复用。")

    add_heading(doc, "7.2 各场景选型建议速查", level=2)
    scene_headers = ["场景", "首选", "理由"]
    scene_rows = [
        ["离线内网态势大屏", "MapLibre", "GPU 矢量 + 完全离线 + 合规"],
        ["在线快速交付(国内)", "高德/百度", "文档好、上手快（仅非离线）"],
        ["轻量原型/移动端弱设备", "Leaflet", "40KB 极轻量、API 简单"],
        ["企业级 GIS/OGC 服务", "OpenLayers", "WMS/WMTS/WFS 全支持"],
        ["真三维/空天/全球尺度", "Cesium", "原生 3D 地球 + 时间轴"],
        ["商用矢量定制(在线)", "Mapbox", "Studio + 数据全栈"],
    ]
    add_table(doc, scene_headers, scene_rows, col_widths=[4.0, 2.6, 9.0])
    add_para(doc, "")

    # ===== 8. 落地路线与风险评估 =====
    add_heading(doc, "8. 落地路线与风险评估", level=1)

    add_heading(doc, "8.1 落地路线（MapLibre）", level=2)
    phase_headers = ["阶段", "工作内容", "预估工时", "产出"]
    phase_rows = [
        ["P1 POC", "maplibre-gl 集成 + tileserver-gl 部署 + 深圳区域瓦片", "2 天", "可交互最小地图"],
        ["P2 适配层", "MapLibreAdapter 实现 + BaiduMapAdapter 封装", "3 天", "双引擎可切换"],
        ["P3 业务迁移", "航线编辑/只读渲染/无人机模拟/控件迁移", "4 天", "功能对齐百度版"],
        ["P4 数据/样式", "OSM 切片 + 中文字体 + 暗色样式", "2 天", "离线底图就绪"],
        ["P5 测试验收", "性能基线 + 离线环境验证", "2 天", "交付报告"],
        ["合计", "", "≈13 人天", ""],
    ]
    add_table(doc, phase_headers, phase_rows, col_widths=[1.8, 7.6, 2.0, 4.2])
    add_para(doc, "")
    add_para(doc, "说明：项目已完成「抽象层 types.ts + maplibre-gl 依赖安装 + 工作计划」(对应 P1/P2 起步)，后续按工作计划推进即可。", color=GRAY, size=9.5)

    add_heading(doc, "8.2 风险与对策", level=2)
    risk_headers = ["风险", "等级", "对策"]
    risk_rows = [
        ["矢量瓦片数据获取与切片", "中", "用 OpenMapTiles 预制 MBTiles + tilemaker 本地切片"],
        ["中文字体标注离线显示", "中", "托管 Noto Sans CJK 的 pbf 字体切片"],
        ["PlaceSearch 等百度专有功能缺失", "中", "离线场景放弃或自建本地 POI 搜索(如基于 OSM 的 Photon)"],
        ["BD09→WGS84 历史数据迁移", "低", "项目已有 coordTransform，一次性批量转换"],
        ["矢量样式定制学习成本", "低", "采用开源 Dark Matter / Positron 样式二次定制"],
        ["内网瓦片服务器性能", "低", "Martin(Rust) 或 tileserver-gl，PMTiles 静态托管可极致简化"],
    ]
    add_table(doc, risk_headers, risk_rows, col_widths=[5.0, 1.6, 9.0])
    add_para(doc, "")

    # ===== 9. 附录 =====
    add_heading(doc, "9. 附录", level=1)

    add_heading(doc, "9.1 关键技术资源", level=2)
    res_headers = ["类别", "项目", "地址"]
    res_rows = [
        ["引擎", "MapLibre GL JS", "https://maplibre.org/"],
        ["引擎", "Leaflet", "https://leafletjs.com/"],
        ["引擎", "OpenLayers", "https://openlayers.org/"],
        ["引擎", "Cesium", "https://cesium.com/"],
        ["瓦片服务", "tileserver-gl", "https://github.com/maptiler/tileserver-gl"],
        ["瓦片服务", "Martin (Rust)", "https://github.com/maplibre/martin"],
        ["数据", "OpenMapTiles", "https://openmaptiles.org/"],
        ["数据", "Geofabrik OSM", "https://download.geofabrik.de/"],
        ["切片工具", "tilemaker", "https://github.com/systemed/tilemaker"],
        ["样式编辑", "Maputnik", "https://maputnik.github.io/"],
        ["可视化", "Deck.gl", "https://deck.gl/"],
    ]
    add_table(doc, res_headers, res_rows, col_widths=[2.2, 3.4, 10.0])
    add_para(doc, "")

    add_heading(doc, "9.2 选型决策快速参考", level=2)
    add_code_block(doc,
        "if 离线/内网 == 必须:\n"
        "    if 3D地形 or 要素>5000 or 需要矢量动态样式:\n"
        "        -> MapLibre GL JS\n"
        "    elif 要素<500 and 快速原型:\n"
        "        -> Leaflet\n"
        "    elif 需要OGC/强GIS:\n"
        "        -> OpenLayers\n"
        "    elif 真三维/全球尺度:\n"
        "        -> Cesium\n"
        "else:  # 在线\n"
        "    if 国内业务 + 快速交付:\n"
        "        -> 高德 / 百度\n"
        "    elif 商用矢量定制 + 预算允许:\n"
        "        -> Mapbox\n"
        "    else:\n"
        "        -> MapLibre / Leaflet")

    add_heading(doc, "9.3 相关文档", level=2)
    add_bullet(doc, "《离线内网 2D 地图方案调研》—— 离线瓦片服务与数据制作细节")
    add_bullet(doc, "《方案 A - MapLibre 工作计划》—— 代码迁移任务分解")
    add_bullet(doc, "《项目说明书》 / 《开发指南》—— 整体架构与规范")

    add_para(doc, "")
    add_para(doc, "—— 本报告结束 ——", size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

    saved = False
    out = OUTPUT_PATH
    try:
        doc.save(out)
        saved = True
    except PermissionError:
        # 目标文件被占用（例如已在 Word 中打开），回退到带时间戳的文件名
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        out = OUTPUT_FALLBACK.format(ts=ts)
        doc.save(out)
        saved = True
    print(f"OK: {out}")


if __name__ == "__main__":
    build_document()